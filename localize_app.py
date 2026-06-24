"""
KRAFTON Japan AIローカライズツール v2
======================================
HQからのKR/ENリリース原稿を日本語に翻訳し、Notion DBに登録する。

【ローカル起動】
  streamlit run localize_app.py
  環境変数: NOTION_TOKEN, ANTHROPIC_KEY, APP_PASSWORD を設定

【Streamlit Cloud デプロイ】
  .streamlit/secrets.toml に以下を記載:
    app_password  = "xxxx"
    NOTION_TOKEN  = "ntn_xxxx"
    ANTHROPIC_KEY = "sk-ant-xxxx"

依存パッケージ:
  pip install streamlit anthropic notion-client python-docx python-dotenv
"""

import os
import io
import re
import json
import streamlit as st
from datetime import date
from notion_client import Client
from anthropic import Anthropic

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ============================================================
# 認証（Streamlit Cloud Simple Auth）
# ============================================================

def _get_secret(key: str, fallback_env: str = "") -> str:
    """st.secrets → 環境変数 の順に値を取得"""
    try:
        return st.secrets[key]
    except Exception:
        return os.environ.get(fallback_env or key, "")


def check_auth() -> bool:
    """パスワード認証ゲート。認証済みなら True を返す。"""
    expected = _get_secret("app_password", "APP_PASSWORD")
    if not expected:
        return True  # パスワード未設定時はスルー（ローカル開発用）

    if st.session_state.get("authenticated"):
        return True

    st.markdown("## 🔐 KRAFTON Japan AIローカライズツール")

    def _verify():
        if st.session_state.get("_pw_input") == expected:
            st.session_state["authenticated"] = True
        else:
            st.session_state["auth_failed"] = True

    st.text_input("パスワード", type="password", key="_pw_input", on_change=_verify)
    if st.session_state.pop("auth_failed", False):
        st.error("パスワードが正しくありません")
    st.stop()


# アプリ起動時に認証チェック
check_auth()

# ============================================================
# 定数
# ============================================================
NOTION_RELEASE_DB_ID   = "7606be38ff494d1f84902f8f82ebfee0"
NOTION_GLOSSARY_DB_ID  = "8959e88bcfe34936a59cc82232215266"
CLAUDE_MODEL           = "claude-sonnet-4-6"

TITLE_TAG_OPTIONS  = ["PUBG", "inZOI", "Subnautica", "KRAFTON", "その他"]
CATEGORY_OPTIONS   = ["コラボ", "アップデート", "大会・イベント", "IR・経営", "採用", "その他"]

# ============================================================
# Notionヘルパー
# ============================================================

@st.cache_data(ttl=300, show_spinner=False)
def fetch_approved_glossary(_token: str) -> list[dict]:
    """approvedステータスの用語をNotionから取得（5分キャッシュ）"""
    notion = Client(auth=_token)
    terms = []
    cursor = None
    while True:
        params = dict(
            database_id=NOTION_GLOSSARY_DB_ID,
            page_size=100,
            filter={"property": "ステータス", "select": {"equals": "approved"}},
        )
        if cursor:
            params["start_cursor"] = cursor
        res = notion.databases.query(**params)
        for page in res["results"]:
            props = page["properties"]
            ja   = _rich_text_value(props.get("日本語表記", {}))
            kr   = _rich_text_value(props.get("KR表記", {}))
            en   = _rich_text_value(props.get("EN表記", {}))
            if ja:
                terms.append({"ja": ja, "kr": kr, "en": en})
        if not res.get("has_more"):
            break
        cursor = res["next_cursor"]
    return terms


def _rich_text_value(prop: dict) -> str:
    """Notionプロパティから文字列を取得"""
    ptype = prop.get("type", "")
    if ptype == "title":
        items = prop.get("title", [])
    elif ptype == "rich_text":
        items = prop.get("rich_text", [])
    else:
        return ""
    return "".join(i.get("plain_text", "") for i in items)


def save_pending_terms(notion: Client, terms: list[dict], source_page_url: str):
    """新出語句をpendingステータスでNotionの用語集DBに保存"""
    for term in terms:
        props = {
            "日本語表記": {"title": [{"text": {"content": term["ja"]}}]},
            "ステータス": {"select": {"name": "pending"}},
        }
        if term.get("kr"):
            props["KR表記"] = {"rich_text": [{"text": {"content": term["kr"]}}]}
        if term.get("en"):
            props["EN表記"] = {"rich_text": [{"text": {"content": term["en"]}}]}
        if source_page_url:
            props["抽出元リリース"] = {"url": source_page_url}
        notion.pages.create(
            parent={"database_id": NOTION_GLOSSARY_DB_ID},
            properties=props,
        )


def save_release_to_notion(
    notion: Client,
    title_ja: str,
    title_original: str,
    body_ja: str,
    body_kr: str,
    body_en: str,
    published_date,
    source_lang: str,
    title_tags: list[str],
    categories: list[str],
) -> tuple[str, str]:
    """ニュースリリースDBにページを作成し (page_id, page_url) を返す"""
    props = {
        "タイトル（日本語）": {"title": [{"text": {"content": title_ja}}]},
        "タイトル（原文）":   {"rich_text": [{"text": {"content": title_original}}]},
        "ステータス":        {"select": {"name": "ai_translated"}},
        "翻訳モデル":        {"rich_text": [{"text": {"content": CLAUDE_MODEL}}]},
    }
    if published_date:
        props["配信日"] = {"date": {"start": published_date.isoformat()}}
    if source_lang in ("KR", "EN"):
        props["原文言語"] = {"select": {"name": source_lang}}
    if title_tags:
        props["タイトルタグ"] = {"multi_select": [{"name": t} for t in title_tags]}
    if categories:
        props["カテゴリ"] = {"multi_select": [{"name": c} for c in categories]}

    page = notion.pages.create(
        parent={"database_id": NOTION_RELEASE_DB_ID},
        properties=props,
    )
    page_id  = page["id"]
    page_url = f"https://app.notion.com/p/{page_id.replace('-', '')}"

    # ページ本文ブロックを構築
    blocks = []

    def add_section(heading: str, text: str):
        if not text.strip():
            return
        blocks.append({
            "object": "block", "type": "heading_2",
            "heading_2": {"rich_text": [{"text": {"content": heading}}]},
        })
        for chunk in [text[i:i+1999] for i in range(0, len(text), 1999)]:
            blocks.append({
                "object": "block", "type": "paragraph",
                "paragraph": {"rich_text": [{"text": {"content": chunk}}]},
            })
        blocks.append({"object": "block", "type": "divider", "divider": {}})

    add_section("日本語訳（AI）", body_ja)
    add_section("原文（English）", body_en)
    add_section("原文（한국어）", body_kr)

    for i in range(0, len(blocks), 50):
        notion.blocks.children.append(block_id=page_id, children=blocks[i:i+50])

    return page_id, page_url


# ============================================================
# Claude翻訳
# ============================================================

def build_glossary_text(terms: list[dict]) -> str:
    if not terms:
        return "（用語集データなし）"
    lines = ["日本語表記 | KR表記 | EN表記"]
    lines.append("-" * 40)
    for t in terms:
        lines.append(f"{t['ja']} | {t.get('kr','―')} | {t.get('en','―')}")
    return "\n".join(lines)


def run_translation(kr_text: str, en_text: str, glossary: list[dict], api_key: str) -> str:
    client = Anthropic(api_key=api_key)

    glossary_text = build_glossary_text(glossary)

    system_prompt = f"""あなたはKRAFTON Japanの公式ニュースリリース翻訳専門家です。
KR/EN原稿を自然で正確な日本語に翻訳します。

【翻訳ルール】
1. プレスリリースとして適切なセミフォーマルな文体
2. 原文の段落構造を維持
3. 数字・日付・URL・製品名は原則そのまま保持
4. 以下の公式用語集に従って表記を統一すること（最優先）

【公式用語集（approved）】
{glossary_text}

【出力形式】
タイトル行（1行目）
空行
本文（段落ごとに空行区切り）
※注釈・コメント不要"""

    if kr_text and en_text:
        user_msg = f"英語を翻訳ベースとし、韓国語でニュアンス・固有名詞を補正してください。\n\n【英語原文】\n{en_text}\n\n【韓国語原文（参考）】\n{kr_text}"
    elif kr_text:
        user_msg = f"以下の韓国語リリースを日本語に翻訳してください。\n\n【韓国語原文】\n{kr_text}"
    else:
        user_msg = f"以下の英語リリースを日本語に翻訳してください。\n\n【英語原文】\n{en_text}"

    res = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": user_msg}],
    )
    return res.content[0].text


def extract_new_terms(translated_ja: str, kr_text: str, en_text: str,
                      approved_glossary: list[dict], api_key: str) -> list[dict]:
    """翻訳結果から新出語句・要確認語句を抽出してリストで返す"""
    client = Anthropic(api_key=api_key)
    approved_ja = {t["ja"] for t in approved_glossary}

    prompt = f"""以下の翻訳済み日本語テキストから、ゲーム固有名詞・タイトル名・イベント名・
組織名・役職名などで「公式用語集に載っていない可能性がある語句」を抽出してください。

【翻訳済み日本語】
{translated_ja[:3000]}

【すでに登録済みの用語（除外してよい）】
{', '.join(list(approved_ja)[:50])}

【出力形式】JSON配列のみ（説明不要）
[
  {{"ja": "日本語表記", "kr": "韓国語表記（不明なら空文字）", "en": "英語表記（不明なら空文字）"}},
  ...
]
抽出対象がなければ [] を返す。"""

    res = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=1024,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = res.content[0].text.strip()
    # JSON部分だけ抽出
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass
    return []


# ============================================================
# ファイル処理
# ============================================================

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"[読み込みエラー: {e}]"


def build_export_docx(title_ja: str, body_ja: str) -> bytes:
    """日本語訳をWord(.docx)にして返す"""
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(title_ja, level=1)
        for para in body_ja.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return b""


# ============================================================
# Streamlit UI
# ============================================================

st.set_page_config(
    page_title="KRAFTON Japan AIローカライズ",
    page_icon="🎮",
    layout="wide",
)

st.title("🎮 KRAFTON Japan AIローカライズツール")
st.caption("HQからのKR/ENリリース原稿を日本語訳し、Notion DBに登録・Exportします。")

# --- サイドバー ---
with st.sidebar:
    st.header("⚙️ 設定")

    # Streamlit Cloud では secrets から自動取得。ローカルは手入力も可。
    _notion_default  = _get_secret("NOTION_TOKEN",  "NOTION_TOKEN")
    _anthropic_default = _get_secret("ANTHROPIC_KEY", "ANTHROPIC_KEY")
    _secrets_loaded = bool(_notion_default and _anthropic_default)

    if _secrets_loaded:
        st.success("🔑 API設定: secrets から読み込み済み")
        notion_token  = _notion_default
        anthropic_key = _anthropic_default
    else:
        notion_token = st.text_input(
            "Notion Token", value=_notion_default,
            type="password", help="https://www.notion.so/my-integrations",
        )
        anthropic_key = st.text_input(
            "Anthropic API Key", value=_anthropic_default,
            type="password", help="https://console.anthropic.com/",
        )

    if notion_token and anthropic_key:
        st.divider()
        with st.spinner("用語集を読み込み中..."):
            glossary = fetch_approved_glossary(notion_token)
        st.success(f"✅ 用語集: {len(glossary)}件 (approved)")
        if st.button("🔄 用語集を再読み込み", use_container_width=True):
            fetch_approved_glossary.clear()
            st.rerun()
        st.caption("[用語集を開く →](https://app.notion.com/p/8959e88bcfe34936a59cc82232215266)")
    else:
        glossary = []
        st.info("API設定を入力してください")

    st.divider()
    st.caption(f"翻訳モデル: `{CLAUDE_MODEL}`")

# --- メインエリア ---
col_input, col_output = st.columns([1, 1], gap="large")

with col_input:
    st.subheader("① 原稿を入力")

    tab_text, tab_file = st.tabs(["📋 テキスト貼り付け", "📁 ファイルアップロード"])

    kr_text = en_text = ""

    with tab_text:
        c1, c2 = st.columns(2)
        with c1:
            kr_text = st.text_area("🇰🇷 韓国語", height=260,
                                   placeholder="韓国語テキストを貼り付け（任意）").strip()
        with c2:
            en_text = st.text_area("🇺🇸 英語", height=260,
                                   placeholder="英語テキストを貼り付け（任意）").strip()

    with tab_file:
        uploaded = st.file_uploader("ドラッグ＆ドロップ or クリック", type=["docx", "txt"])
        if uploaded:
            raw = (extract_text_from_docx(uploaded.read())
                   if uploaded.name.endswith(".docx")
                   else uploaded.read().decode("utf-8", errors="replace"))
            file_lang = st.radio("言語", ["KR", "EN", "KR+EN混在"],
                                 horizontal=True, key="file_lang")
            st.text_area("抽出テキスト（確認用）", value=raw, height=180, disabled=True)
            if file_lang == "KR":
                kr_text = raw
            elif file_lang == "EN":
                en_text = raw
            else:
                mid = len(raw) // 2
                kr_text, en_text = raw[:mid], raw[mid:]
                st.info("混在の場合はテキスト貼り付けタブで分割入力する方が正確です。")

    detected = []
    if kr_text: detected.append("🇰🇷 KR")
    if en_text: detected.append("🇺🇸 EN")
    if detected:
        st.success(f"入力検知: {' + '.join(detected)}")
    else:
        st.info("原文を入力してください")

    st.divider()
    st.subheader("② メタデータ")

    title_original = st.text_input("原文タイトル", placeholder="KR or EN のタイトル")
    pub_date = st.date_input("配信予定日", value=date.today())
    c1, c2 = st.columns(2)
    with c1: selected_tags = st.multiselect("タイトルタグ", TITLE_TAG_OPTIONS)
    with c2: selected_cats = st.multiselect("カテゴリ", CATEGORY_OPTIONS)

    st.divider()
    can_go = bool((kr_text or en_text) and notion_token and anthropic_key)
    if st.button("🤖 AI翻訳を実行", disabled=not can_go,
                 use_container_width=True, type="primary"):
        with st.spinner("Claude APIで翻訳中..."):
            try:
                result = run_translation(kr_text, en_text, glossary, anthropic_key)
                st.session_state.update({
                    "translated": result,
                    "translate_done": True,
                    "notion_saved": False,
                    "saved_url": "",
                    "new_terms": [],
                })
            except Exception as e:
                st.error(f"翻訳エラー: {e}")

    if not can_go:
        st.caption("原文 + API設定が必要です")


with col_output:
    st.subheader("③ 翻訳確認・編集")

    if st.session_state.get("translate_done"):
        body_edited = st.text_area(
            "日本語訳（編集可）",
            value=st.session_state.get("translated", ""),
            height=320, key="body_edit",
        )
        lines = body_edited.strip().split("\n")
        title_ja = st.text_input(
            "日本語タイトル（編集可）",
            value=lines[0].strip() if lines else "",
            key="title_ja_edit",
        )

        st.divider()

        # --- Notionへ登録 ---
        c1, c2 = st.columns(2)
        with c1:
            if st.button("✅ Notionに登録", use_container_width=True, type="primary"):
                if not title_ja.strip():
                    st.error("タイトルを入力してください")
                else:
                    src_lang = "KR" if kr_text and not en_text else \
                               ("EN" if en_text and not kr_text else "KR")
                    with st.spinner("Notionに登録中..."):
                        try:
                            notion = Client(auth=notion_token)
                            pid, purl = save_release_to_notion(
                                notion, title_ja, title_original or title_ja,
                                body_edited, kr_text, en_text,
                                pub_date, src_lang, selected_tags, selected_cats,
                            )
                            st.session_state["notion_saved"] = True
                            st.session_state["saved_url"] = purl
                            st.session_state["saved_pid"] = pid

                            # 新出語句の抽出
                            with st.spinner("新出語句を抽出中..."):
                                new_terms = extract_new_terms(
                                    body_edited, kr_text, en_text, glossary, anthropic_key
                                )
                                if new_terms:
                                    save_pending_terms(notion, new_terms, purl)
                                st.session_state["new_terms"] = new_terms
                        except Exception as e:
                            st.error(f"登録エラー: {e}")

        with c2:
            # Export（Wordダウンロード）
            docx_bytes = build_export_docx(
                st.session_state.get("title_ja_edit", ""),
                st.session_state.get("body_edit", ""),
            )
            if docx_bytes:
                safe_title = re.sub(r'[\\/:*?"<>|]', "_",
                                    st.session_state.get("title_ja_edit", "release"))[:40]
                st.download_button(
                    "📄 Word出力 (.docx)",
                    data=docx_bytes,
                    file_name=f"{safe_title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            else:
                st.caption("python-docx未インストール時はWord出力不可")

        # --- 登録後の表示 ---
        if st.session_state.get("notion_saved"):
            st.success("✅ Notionに登録しました（ステータス: ai_translated）")
            st.link_button("Notionで確認・校了する →",
                           st.session_state.get("saved_url", ""))

            new_terms = st.session_state.get("new_terms", [])
            if new_terms:
                st.warning(f"📝 新出語句 {len(new_terms)}件を用語集にpendingで追記しました")
                with st.expander("追記された用語を確認する"):
                    for t in new_terms:
                        st.write(f"- **{t['ja']}**"
                                 + (f"（KR: {t['kr']}）" if t.get("kr") else "")
                                 + (f"（EN: {t['en']}）" if t.get("en") else ""))
                st.link_button("用語集で承認する（Admin）→",
                               "https://app.notion.com/p/8959e88bcfe34936a59cc82232215266")
            else:
                st.info("新出語句は検出されませんでした")

    else:
        st.info("原文を入力して「AI翻訳を実行」を押すと、ここに結果が表示されます。")
        st.markdown("""
**全体の流れ:**

| ステップ | 操作 | ステータス |
|----------|------|-----------|
| 1 | 原稿インポート | — |
| 2 | AI翻訳 | — |
| 3 | 人間が確認・編集 | — |
| 4 | Notionに登録 | `ai_translated` |
| 5 | Word出力して配布 | — |
| 6 | Notionで最終校了 | `approved` |
| 7 | 新出語句をAdmin承認 | 用語集に反映 |
""")
